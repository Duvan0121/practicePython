import pyautogui
from docx import Document
from docx.shared import Inches
import os
from docx2pdf import convert

# Creamos el documento de Word globalmente
doc = Document()
base_path = 'C:/Users/varga/PycharmProjects/automatizacionPython/Evidencias'
evidence_folder = None


def get_next_folder_number():
    # Obtener la lista de carpetas dentro de la ruta especificada
    folders = [name for name in os.listdir(base_path) if os.path.isdir(os.path.join(base_path, name))]
    # Obtener el número más alto dentro de los nombres de las carpetas
    max_number = max([int(folder) for folder in folders]) if folders else 0
    # Retornar el siguiente número disponible
    return max_number + 1

def create_folder(folder_name):
    folder_path = os.path.join(base_path, folder_name)
    os.makedirs(folder_path)
    return folder_path

def create_evidence_folder():
    global evidence_folder
    if evidence_folder is None:
        # Si no se ha creado la carpeta antes, creamos una nueva
        next_folder_number = get_next_folder_number()
        evidence_folder = create_folder(str(next_folder_number))
    return evidence_folder

def create_pdf_from_word(word_file_path, pdf_file_path):
    convert(word_file_path, pdf_file_path)

def evidence(heading_text):
    global doc  # Accedemos a la variable global doc

    # Obtener la carpeta de evidencia
    evidence_folder = create_evidence_folder()

    # Capturar la pantalla
    screenshot = pyautogui.screenshot()

    # Guardar la captura de pantalla en un archivo temporal dentro de la carpeta
    screenshot_path = os.path.join(evidence_folder, "temp_screenshot.png")
    screenshot.save(screenshot_path)

    # Agregar el encabezado al documento de Word
    doc.add_heading(heading_text, 0)

    # Agregar la captura de pantalla al documento de Word
    doc.add_picture(screenshot_path, width=Inches(6))
    doc.add_paragraph("")  # Añadir un párrafo en blanco para separar las imágenes

    # Eliminar la captura de pantalla temporal
    os.remove(screenshot_path)

    # Guardar el documento de Word dentro de la carpeta de evidencia
    doc_path = os.path.join(evidence_folder, "EvidenciaEjecucion.docx")
    doc.save(doc_path)
    print(f"Evidencia de: '{heading_text}' guardada")

    # Convertir el archivo Word a PDF y guardarlo en la misma carpeta
    #create_pdf_from_word(doc_path, os.path.join(evidence_folder, "Evidencias.pdf"))

